"""Анализ прикреплённых файлов заказа (PDF, DOCX, изображения) — извлечение текста и суммаризация."""

import base64
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from src.ai_client import chat_completion, chat_completion_vision
from src.config import settings

logger = logging.getLogger(__name__)

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".heic", ".bmp", ".gif", ".webp"}

VISION_PROMPT = (
    "Извлеки ВЕСЬ текст с этого изображения. Включи всё: условия задачи, "
    "числа, формулы, таблицы, подписи. Если есть рукописный текст — распознай его. "
    "Верни только извлечённый текст, без комментариев."
)


def is_image_file(file_path: Path) -> bool:
    """Проверить, является ли файл изображением."""
    return file_path.suffix.lower() in IMAGE_EXTENSIONS


def _get_mime_type(file_path: Path) -> str:
    """Определить MIME-тип изображения."""
    suffix = file_path.suffix.lower()
    mime_map = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".gif": "image/gif",
        ".webp": "image/webp",
        ".bmp": "image/bmp",
        ".heic": "image/heic",
    }
    return mime_map.get(suffix, "image/jpeg")


def extract_text_from_pdf(file_path: Path) -> str:
    """Извлечь текст из PDF через PyMuPDF."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(file_path))
        text_parts = []
        for page in doc:
            text_parts.append(page.get_text())
        doc.close()
        return "\n".join(text_parts)
    except Exception as e:
        logger.error("Ошибка извлечения текста из PDF %s: %s", file_path, e)
        return ""


def extract_text_from_docx(file_path: Path) -> str:
    """Извлечь текст из DOCX через python-docx."""
    try:
        from docx import Document
        doc = Document(str(file_path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error("Ошибка извлечения текста из DOCX %s: %s", file_path, e)
        return ""


def extract_text(file_path: Path) -> str:
    """Извлечь текст из файла (определяет тип по расширению).

    Для изображений возвращает пустую строку — используйте extract_text_from_image().
    """
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    elif suffix == ".txt":
        try:
            return file_path.read_text(encoding="utf-8")
        except Exception:
            return file_path.read_text(encoding="cp1251", errors="ignore")
    elif suffix in IMAGE_EXTENSIONS:
        return ""  # Изображения обрабатываются через vision API
    else:
        logger.warning("Неподдерживаемый формат файла: %s", suffix)
        return ""


# ---------------------------------------------------------------------------
# Vision: извлечение текста из изображений через GPT-4o
# ---------------------------------------------------------------------------

async def extract_text_from_image(file_path: Path) -> dict:
    """Извлечь текст из изображения через GPT-4o vision.

    Returns:
        {"text": str, "input_tokens": int, "output_tokens": int, "cost_usd": float}
    """
    try:
        image_data = base64.b64encode(file_path.read_bytes()).decode()
        mime = _get_mime_type(file_path)

        result = await chat_completion_vision(
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": VISION_PROMPT},
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{image_data}"}},
                ],
            }],
            max_tokens=2048,
        )

        text = result["content"].strip()
        logger.info(
            "Распознано изображение %s: %s",
            file_path.name, text[:100] + "..." if len(text) > 100 else text,
        )

        return {
            "text": text,
            "input_tokens": result["input_tokens"],
            "output_tokens": result["output_tokens"],
            "cost_usd": result["cost_usd"],
        }
    except Exception as e:
        logger.error("Ошибка распознавания изображения %s: %s", file_path, e)
        return {"text": "", "input_tokens": 0, "output_tokens": 0, "cost_usd": 0}


async def extract_text_from_image_bytes(image_bytes: bytes, name: str = "image") -> dict:
    """Извлечь текст из изображения в виде байтов (для встроенных изображений из PDF/DOCX).

    Returns:
        {"text": str, "input_tokens": int, "output_tokens": int, "cost_usd": float}
    """
    try:
        image_data = base64.b64encode(image_bytes).decode()

        result = await chat_completion_vision(
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": VISION_PROMPT},
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{image_data}"}},
                ],
            }],
            max_tokens=2048,
        )

        text = result["content"].strip()
        logger.info(
            "Распознано встроенное изображение (%s): %s",
            name, text[:100] + "..." if len(text) > 100 else text,
        )

        return {
            "text": text,
            "input_tokens": result["input_tokens"],
            "output_tokens": result["output_tokens"],
            "cost_usd": result["cost_usd"],
        }
    except Exception as e:
        logger.error("Ошибка распознавания встроенного изображения (%s): %s", name, e)
        return {"text": "", "input_tokens": 0, "output_tokens": 0, "cost_usd": 0}


def extract_image_pages_from_pdf(file_path: Path) -> list[bytes]:
    """Извлечь изображения из PDF-страниц, которые не содержат текста.

    Возвращает список PNG-байтов для каждой страницы-изображения.
    """
    result = []
    try:
        import fitz
        doc = fitz.open(str(file_path))
        for i, page in enumerate(doc):
            text = page.get_text().strip()
            if not text:
                # Страница без текста — вероятно, скан/изображение
                pix = page.get_pixmap(dpi=200)
                result.append(pix.tobytes("png"))
                logger.info("PDF %s: страница %d — изображение (без текста)", file_path.name, i + 1)
        doc.close()
    except Exception as e:
        logger.error("Ошибка извлечения изображений из PDF %s: %s", file_path, e)
    return result


def extract_images_from_docx(file_path: Path) -> list[bytes]:
    """Извлечь встроенные изображения из DOCX файла."""
    result = []
    try:
        from docx import Document
        doc = Document(str(file_path))
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_bytes = rel.target_part.blob
                    result.append(image_bytes)
                except Exception:
                    continue
        if result:
            logger.info("DOCX %s: найдено %d встроенных изображений", file_path.name, len(result))
    except Exception as e:
        logger.error("Ошибка извлечения изображений из DOCX %s: %s", file_path, e)
    return result


# ---------------------------------------------------------------------------
# Комплексное извлечение контента (текст + vision)
# ---------------------------------------------------------------------------

@dataclass
class ContentResult:
    """Результат извлечения контента из файлов."""
    text: str = ""
    vision_texts: list[str] = field(default_factory=list)
    total_input_tokens: int = 0
    total_output_tokens: int = 0
    total_cost_usd: float = 0.0

    @property
    def all_text(self) -> str:
        """Весь текст: из файлов + из vision."""
        parts = [self.text] if self.text else []
        parts.extend(self.vision_texts)
        return "\n\n".join(parts)


async def extract_all_content(file_paths: list[Path]) -> ContentResult:
    """Извлечь весь контент из файлов: текст + изображения через vision.

    Returns:
        ContentResult с текстом и метаданными о vision-вызовах.
    """
    result = ContentResult()
    text_parts = []

    for fp in file_paths:
        suffix = fp.suffix.lower()

        if is_image_file(fp):
            # Отдельное изображение → vision
            vision_result = await extract_text_from_image(fp)
            if vision_result["text"]:
                result.vision_texts.append(f"[Изображение {fp.name}]\n{vision_result['text']}")
            result.total_input_tokens += vision_result["input_tokens"]
            result.total_output_tokens += vision_result["output_tokens"]
            result.total_cost_usd += vision_result["cost_usd"]

        elif suffix == ".pdf":
            # PDF: текст + изображения со страниц без текста
            text = extract_text_from_pdf(fp)
            if text.strip():
                text_parts.append(f"--- Файл: {fp.name} ---\n{text}")

            # Страницы без текста → vision
            image_pages = extract_image_pages_from_pdf(fp)
            for i, img_bytes in enumerate(image_pages):
                vision_result = await extract_text_from_image_bytes(
                    img_bytes, name=f"{fp.name} стр.{i+1}"
                )
                if vision_result["text"]:
                    result.vision_texts.append(
                        f"[{fp.name} стр.{i+1} (скан)]\n{vision_result['text']}"
                    )
                result.total_input_tokens += vision_result["input_tokens"]
                result.total_output_tokens += vision_result["output_tokens"]
                result.total_cost_usd += vision_result["cost_usd"]

        elif suffix in (".docx", ".doc"):
            # DOCX: текст + встроенные изображения
            text = extract_text_from_docx(fp)
            if text.strip():
                text_parts.append(f"--- Файл: {fp.name} ---\n{text}")

            # Встроенные изображения → vision
            images = extract_images_from_docx(fp)
            for i, img_bytes in enumerate(images):
                vision_result = await extract_text_from_image_bytes(
                    img_bytes, name=f"{fp.name} img#{i+1}"
                )
                if vision_result["text"]:
                    result.vision_texts.append(
                        f"[{fp.name} изобр.{i+1}]\n{vision_result['text']}"
                    )
                result.total_input_tokens += vision_result["input_tokens"]
                result.total_output_tokens += vision_result["output_tokens"]
                result.total_cost_usd += vision_result["cost_usd"]

        else:
            # Прочие текстовые файлы
            text = extract_text(fp)
            if text.strip():
                text_parts.append(f"--- Файл: {fp.name} ---\n{text}")

    result.text = "\n\n".join(text_parts)
    return result


# ---------------------------------------------------------------------------
# Суммаризация (оригинальная функция, обновлена для vision)
# ---------------------------------------------------------------------------

async def summarize_files(file_paths: list[Path]) -> Optional[dict]:
    """Извлечь и суммаризировать содержимое прикреплённых файлов.

    Returns:
        {
            "summary": str,
            "requirements": str,
            "structure": str,
            "volume": str,
            "raw_text": str (усечённый),
            "input_tokens": int,
            "output_tokens": int,
            "cost_usd": float,
        }
    """
    # Используем extract_all_content для получения текста (включая vision)
    content_result = await extract_all_content(file_paths)
    combined = content_result.all_text

    if not combined.strip():
        return None

    # Ограничиваем текст для API
    if len(combined) > 12000:
        combined = combined[:12000] + "\n... (текст обрезан)"

    result = await chat_completion(
        messages=[
            {
                "role": "system",
                "content": (
                    "Ты анализируешь методические указания и приложенные файлы к заказу. "
                    "Извлеки ключевую информацию для написания работы.\n\n"
                    "Ответь структурированно:\n"
                    "1. КРАТКОЕ СОДЕРЖАНИЕ: о чём методичка/файл\n"
                    "2. ТРЕБОВАНИЯ К ОФОРМЛЕНИЮ: шрифт, интервал, поля, нумерация и т.д.\n"
                    "3. СТРУКТУРА РАБОТЫ: какие разделы/главы нужны\n"
                    "4. ОБЪЁМ: сколько страниц, символов или слов"
                ),
            },
            {"role": "user", "content": f"Проанализируй следующие файлы:\n\n{combined}"},
        ],
        model=settings.openai_model_fast,
        temperature=0.2,
        max_tokens=1024,
    )

    content = result["content"]

    # Разбираем ответ на секции
    sections = content.split("\n")
    current_section = "summary"
    section_lines: dict[str, list[str]] = {"summary": [], "requirements": [], "structure": [], "volume": []}

    for line in sections:
        lower = line.lower()
        if "требовани" in lower or "оформлени" in lower:
            current_section = "requirements"
        elif "структур" in lower or "раздел" in lower or "глав" in lower:
            current_section = "structure"
        elif "объём" in lower or "объем" in lower or "страниц" in lower:
            current_section = "volume"
        section_lines[current_section].append(line)

    # Суммируем токены: vision + summarization
    total_in = content_result.total_input_tokens + result["input_tokens"]
    total_out = content_result.total_output_tokens + result["output_tokens"]
    total_cost = content_result.total_cost_usd + result["cost_usd"]

    return {
        "summary": "\n".join(section_lines["summary"]).strip(),
        "requirements": "\n".join(section_lines["requirements"]).strip(),
        "structure": "\n".join(section_lines["structure"]).strip(),
        "volume": "\n".join(section_lines["volume"]).strip(),
        "raw_text": combined[:5000],
        "input_tokens": total_in,
        "output_tokens": total_out,
        "cost_usd": total_cost,
    }
